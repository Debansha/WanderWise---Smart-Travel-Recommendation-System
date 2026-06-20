import os
import time
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

pages = [
    {"file": "index.html", "title": "1. Home / Landing Page", "desc": "The central hub of WanderWise, featuring a breathtaking particle hero section, an auto-scrolling testimonials slider, and an interactive features grid. It showcases the premium glassmorphism design and provides quick access to core functionalities."},
    {"file": "destinations.html", "title": "2. Destinations Explorer", "desc": "A dynamic gallery of global travel destinations. Includes category and continent filters, search functionality, and aesthetic tilt-hover effect cards displaying pricing and ratings for curated travel spots."},
    {"file": "planner.html", "title": "3. AI Trip Planner", "desc": "A multi-step interactive wizard that gathers user preferences (interests, budget, and weather) and processes them to suggest highly personalized travel itineraries based on logic and filtering."},
    {"file": "budget.html", "title": "4. Budget Calculator", "desc": "A smart financial tool offering interactive sliders for budget allocation across accommodation, transport, food, and activities. It visualizes the breakdown using an intricate CSS/Canvas-based donut chart and includes a real-time currency converter."},
    {"file": "weather.html", "title": "5. Weather Dashboard", "desc": "Provides average climate insights for select destinations. It includes interactive temperature charts configured dynamically with JavaScript, alongside filterable best-season recommendations."},
    {"file": "culture.html", "title": "6. Culture Guide", "desc": "An immersive guide detailing local cuisines, festivals, do's & don'ts, and essential local phrases. Organized efficiently using a responsive tabbed interface allowing comparison among various travel hotspots."},
    {"file": "itinerary.html", "title": "7. Itinerary Builder", "desc": "A tool where users can select a destination and duration to auto-load a template, manually add or modify activities dynamically, and generate an exportable travel itinerary text file."},
    {"file": "blog.html", "title": "8. Travel Blog", "desc": "A magazine-style reading interface with category filtering, reading time estimates, and featured travel articles designed heavily utilizing modern typography and card layouts."},
    {"file": "gallery.html", "title": "9. Photo Gallery", "desc": "A visually striking masonry grid layout for global travel photography. It features filtering by theme (nature, cities, beaches) and a fully functional custom javascript lightbox for enlarged image viewing."},
    {"file": "contact.html", "title": "10. Contact & About Us", "desc": "Features an animated feedback form module, an interactive team introduction grid, business hours, and a collapsible FAQ accordion to address user inquiries."},
    {"file": "tips.html", "title": "11. Travel Tips & Hacks", "desc": "A utility page containing an interactive, progress-tracked packing checklist, categorized overarching safety guidelines, and visual travel hack cards."}
]

base_dir = r"c:\Users\deban\OneDrive\Documents\CODING\TravelAssistant"
os.chdir(base_dir)

if not os.path.exists("shots"):
    os.makedirs("shots")

print("Initializing Word Document...")
doc = Document()

# Cover Page
title = doc.add_heading('WANDERWISE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Bilingual Travel Assistant Platform\n\n', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Project Documentation & Source Code\n\n\n').bold = True

p_details = doc.add_paragraph()
p_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_details.add_run('Submitted By:\n').bold = True
p_details.add_run('Name: [Your Name Here]\n')
p_details.add_run('Registration Number: [Your Reg Number]\n')
p_details.add_run('Course/Section: [Your Course]')
doc.add_page_break()

# Index Page
doc.add_heading('Table of Contents', 1)
index_p = doc.add_paragraph()
for item in pages:
    index_p.add_run(item["title"] + '\n')
index_p.add_run('\nShared Application Files:\n')
index_p.add_run('12. Main Stylesheet (style.css)\n')
index_p.add_run('13. Main Logic (main.js)\n')
index_p.add_run('14. Bilingual Dictionary (translations.js)\n')
doc.add_page_break()

# Pages Content
with sync_playwright() as p:
    print("Launching browser...")
    browser = p.chromium.launch(headless=True)
    page = browser.new_page(viewport={"width": 1280, "height": 800})
    
    for item in pages:
        print(f"Processing {item['title']}...")
        file_path = os.path.join(base_dir, item["file"])
        file_url = f"file:///{file_path.replace(chr(92), '/')}"
        
        print(f"  Taking screenshot of {item['file']}...")
        page.goto(file_url, wait_until="networkidle")
        time.sleep(1.5)
        shot_path = f"shots/{item['file']}.png"
        
        # full_page=False takes a viewport-sized screenshot suitable for a Word doc page
        page.screenshot(path=shot_path, full_page=False)
        
        print(f"  Reading code for {item['file']}...")
        with open(file_path, "r", encoding="utf-8") as f:
            code = f.read()
            
        # Write to document matching requirement layout
        doc.add_heading(item["title"], level=1)
        
        doc.add_picture(shot_path, width=Inches(6.0))
        
        doc.add_heading('Description', level=2)
        doc.add_paragraph(item["desc"])
        
        doc.add_heading('Source Code', level=2)
        code_para = doc.add_paragraph()
        run = code_para.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        
        doc.add_page_break()
        
    assets = [
        {"file": "css/style.css", "title": "12. Core Stylesheet (style.css)", "desc": "Defines the premium dark glassmorphism aesthetic, color variables, layout grids, and animations across all pages."},
        {"file": "js/main.js", "title": "13. Application Logic (main.js)", "desc": "Handles core interactivity including language toggling, navigation menu state, particle canvas rendering, scroll reveals, and the custom cursor."},
        {"file": "js/translations.js", "title": "14. Translation Dictionary (translations.js)", "desc": "A robust JSON-based translation object mapping keys to English and Hindi strings, enabling the bilingual capabilities of the platform."}
    ]
    
    for asset in assets:
        print(f"Processing {asset['title']}...")
        file_path = os.path.join(base_dir, asset["file"])
        with open(file_path, "r", encoding="utf-8") as f:
            code = f.read()
            
        doc.add_heading(asset["title"], level=1)
        doc.add_heading('Description', level=2)
        doc.add_paragraph(asset["desc"])
        
        doc.add_heading('Source Code', level=2)
        code_para = doc.add_paragraph()
        run = code_para.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        
        doc.add_page_break()

    browser.close()

doc_path = "WanderWise_Final_Submission.docx"
doc.save(doc_path)
print(f"Documentation generated successfully: {doc_path}")
